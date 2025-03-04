import os
import docx
import pdfplumber
import google.generativeai as genai
from deep_translator import GoogleTranslator

# Configure Google Gemini AI
genai.configure(api_key="AIzaSyBRbOwTCb015gqn4ii2WhB1-ah5RFrttfQ")

# Function to extract text from TXT
def extract_text_from_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()

# Function to extract text from PDF
def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])

# Function to extract text from DOCX
def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# Function to process file and translate
def process_file(filename, target_lang):
    file_path = os.path.join("documents", filename)

    if not os.path.exists(file_path):
        return None, "File not found."

    # Determine file type
    if filename.endswith(".txt"):
        text = extract_text_from_txt(file_path)
    elif filename.endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    elif filename.endswith(".docx"):
        text = extract_text_from_docx(file_path)
    else:
        return None, "Unsupported file format"

    if not text:
        return None, "Empty document or unreadable content"

    # Translate text
    translated_text = GoogleTranslator(source="auto", target=target_lang).translate(text)

    # Generate expanded content using Gemini AI
    prompt = f"Expand and enhance the following text:\n\n{translated_text}"
    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(prompt)
    expanded_text = response.text if response.text else "No additional content generated."

    # Save the translated content
    translated_filename = f"translated_{target_lang}_{filename}"
    translated_path = os.path.join("translated_documents", translated_filename)

    os.makedirs("translated_documents", exist_ok=True)

    if filename.endswith(".txt"):
        with open(translated_path, "w", encoding="utf-8") as f:
            f.write(expanded_text)
    elif filename.endswith(".docx"):
        doc = docx.Document()
        doc.add_paragraph(expanded_text)
        doc.save(translated_path)
    elif filename.endswith(".pdf"):
        with open(translated_path, "w", encoding="utf-8") as f:
            f.write(expanded_text)  # Save PDF as text for now (or use a PDF writer library)

    return translated_path, None
